import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import pandas as pd
import os
import win32com.client as win32
import logging
import re
from openpyxl import Workbook, load_workbook
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from collections import Counter
from tkinter import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
background = "#06283D"


background = "#06283D"

# Configuration de la journalisation
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')

# Variables globales pour stocker les données
base_de_donnees_manifeste = None
fichier_xlsb = None
root = None  # Variable globale pour la fenêtre principale Tkinter
dernieres_valeurs = []  # Tableau pour stocker les dernières valeurs de la colonne A
wb = None  # Variable globale pour le workbook Excel
ws = None  # Variable globale pour la feuille Excel

# Charger le fichier Excel contenant les données de localisation
df_localisation = pd.read_excel('./data/liste des villes/localisation.xlsx')

# Vérifier les noms de colonnes du DataFrame
logging.debug(f"Colonnes du DataFrame localisation : {df_localisation.columns.tolist()}")

# Fonction pour normaliser le texte
def normalize_text(text):
    text = re.sub(r"['’`]", "", text)  # Enlever les apostrophes et caractères similaires
    text = re.sub(r'\s+', ' ', text)  # Remplacer les espaces multiples par un seul espace
    return text.strip().lower()  # Supprimer les espaces en trop et mettre en minuscule

# Fonction pour trouver les correspondances dans le DataFrame en lisant l'adresse
def trouver_info_adresse(adresse, df):
    mots = normalize_text(adresse).split()
    fokontany_candidate = ''

    # Construire un pattern regex à partir des fokontany, kaomina et distrika
    fokontany_pattern = '|'.join(map(re.escape, df['fokontany'].dropna().unique()))
    kaomina_pattern = '|'.join(map(re.escape, df['Kaomina'].dropna().unique()))
    distrika_pattern = '|'.join(map(re.escape, df['Distrika'].dropna().unique()))
    
    # Rechercher les fokontany, kaomina et distrika dans l'adresse
    fokontany_match = re.search(fokontany_pattern, adresse, re.IGNORECASE)
    kaomina_match = re.search(kaomina_pattern, adresse, re.IGNORECASE)
    distrika_match = re.search(distrika_pattern, adresse, re.IGNORECASE)
    
    # Vérifier les correspondances et retourner les informations
    if fokontany_match:
        fokontany_candidate = fokontany_match.group()
        match = df[df['fokontany'].str.contains(fokontany_candidate, case=False, na=False)]
    elif kaomina_match:
        kaomina_candidate = kaomina_match.group()
        match = df[df['Kaomina'].str.contains(kaomina_candidate, case=False, na=False)]
    elif distrika_match:
        distrika_candidate = distrika_match.group()
        match = df[df['Distrika'].str.contains(distrika_candidate, case=False, na=False)]
    else:
        return None, None, None, None, None, None
    
    if not match.empty:
        kaomina = match['Kaomina'].values[0]
        distrika = match['Distrika'].values[0]
        secteur = match['Secteur'].values[0]
        region = match['Region'].values[0]
        province = match['Province'].values[0]
        return fokontany_candidate, kaomina, distrika, secteur, region, province
    
    return None, None, None, None, None, None

# Fonction pour afficher les informations correspondantes dans une fenêtre popup
def afficher_informations_correspondantes(informations):
    logging.debug("Affichage des informations correspondantes dans une fenêtre popup.")
    popup = tk.Toplevel()
    popup.title("Informations Correspondantes")
    popup.config(bg="black")

    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()
    popup_width = screen_width // 2
    popup_height = screen_height

    popup.geometry(f"{popup_width}x{popup_height}+{screen_width // 2}+0")

    # Récupérer les informations nécessaires
    awb = informations['AWB']
    consignee_name = informations['ConsigneeName']
    consignee_address = informations['ConsigneeAddress']
    dest_city = informations['DestCity']

    # Utiliser la nouvelle méthode pour trouver les informations d'adresse
    fokontany, kaomina, distrika, secteur, region, province = trouver_info_adresse(consignee_address, df_localisation)

    # Affichage des informations dans des labels
    labels = []
    label_awb = ttk.Label(popup, text=f"AWB: {awb}", font="Arial 20 bold", foreground='white', background='black')
    label_awb.grid(row=0, column=0, padx=20, pady=30, sticky="w")
    labels.append(label_awb)

    label_consignee_name = ttk.Label(popup, text=f"Name: {consignee_name}", font="Arial 20 bold", foreground='white', background='black')
    label_consignee_name.grid(row=3, column=0, padx=20, pady=30, sticky="w")
    labels.append(label_consignee_name)
    
    label_consignee_secteur = ttk.Label(popup, text=f"Secteur: {secteur}", font="Arial 30 bold", foreground='orange', background='black')
    label_consignee_secteur.grid(row=4, column=0, padx=20, pady=30, sticky="w")
    labels.append(label_consignee_secteur)

    label_consignee_address = ttk.Label(popup, text=f"Address: {consignee_address}", font="Arial 16 bold", foreground='green', background='black')
    label_consignee_address.grid(row=5, column=0, padx=20, pady=30, sticky="w")
    labels.append(label_consignee_address)

    label_consignee_kaomina = ttk.Label(popup, text=f"Kaomina: {kaomina}", font="Arial 18 bold", foreground='white', background='black')
    label_consignee_kaomina.grid(row=6, column=0, padx=20, pady=30, sticky="w")
    labels.append(label_consignee_kaomina)

    label_dest_city = ttk.Label(popup, text=f"Destination City: {dest_city}", font="Arial 16 bold", foreground='white', background='black')
    label_dest_city.grid(row=7, column=0, padx=20, pady=30, sticky="w")
    labels.append(label_dest_city)

    popup.attributes("-topmost", True)  # Mettre la fenêtre en avant-plan
    popup.update_idletasks()
    popup.after(9000, popup.destroy)  # Fermer le popup après 9 secondes

# Fonction pour choisir le fichier Excel contenant la base de données manifeste
def choisir_fichier_excel():
    logging.debug("Ouverture de la boîte de dialogue pour choisir le fichier Excel.")
    root = tk.Tk()
    root.withdraw()  # Cacher la fenêtre principale tkinter
    file_path = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx;*.xls")])
    logging.debug(f"Fichier Excel sélectionné : {file_path}")
    return file_path

# Fonction pour choisir le fichier XLSB
def choisir_fichier_xlsb():
    logging.debug("Ouverture de la boîte de dialogue pour choisir le fichier XLSB.")
    root = tk.Tk()
    root.withdraw()  # Cacher la fenêtre principale tkinter
    file_path = filedialog.askopenfilename(filetypes=[("XLSB Files", "*.xlsb")])
    logging.debug(f"Fichier XLSB sélectionné : {file_path}")
    return file_path

# Fonction pour charger la base de données manifeste à l'aide de pandas
def charger_base_de_donnees_manifeste(file_excel):
    global base_de_donnees_manifeste
    try:
        logging.debug(f"Chargement de la base de données manifeste depuis {file_excel}")
        base_de_donnees_manifeste = pd.read_excel(file_excel)
        return base_de_donnees_manifeste
    except Exception as e:
        logging.error(f"Erreur lors du chargement de la base de données manifeste : {e}")
        return None

# Fonction pour charger et afficher la base de données manifeste
def charger_et_afficher_base_de_donnees():
    global base_de_donnees_manifeste
    global fichier_xlsb

    fichier_excel = choisir_fichier_excel()

    if fichier_excel:
        base_de_donnees_manifeste = charger_base_de_donnees_manifeste(fichier_excel)

        if base_de_donnees_manifeste is not None:
            afficher_donnees_treeview(treeview, base_de_donnees_manifeste)

            fichier_xlsb = choisir_fichier_xlsb()

            if fichier_xlsb:
                ouvrir_fichier_xlsb()

                root.update_idletasks()
                root_width = root.winfo_screenwidth() // 2
                root_height = root.winfo_screenheight()
                root.geometry(f"{root_width}x{root_height}+0+0")

                verifier_changements()  # Lancer la vérification périodique

                root.mainloop()
            else:
                messagebox.showerror("Erreur", "Aucun fichier XLSB sélectionné.")
        else:
            messagebox.showerror("Erreur", "Impossible de charger la base de données manifeste.")
    else:
        messagebox.showerror("Erreur", "Aucun fichier Excel sélectionné.")

# Fonction pour ouvrir un fichier XLSB
def ouvrir_fichier_xlsb():
    global fichier_xlsb, wb, ws
    if fichier_xlsb:
        try:
            logging.debug(f"Ouverture du fichier XLSB : {fichier_xlsb}")
            excel = win32.Dispatch('Excel.Application')
            wb = excel.Workbooks.Open(fichier_xlsb)
            ws = wb.Worksheets(1)
            excel.Visible = True  # Laisser le fichier ouvert et visible
        except Exception as e:
            logging.error(f"Erreur lors de l'ouverture du fichier XLSB : {e}")
    else:
        logging.debug("Aucun fichier XLSB sélectionné.")

# Fonction pour afficher les données dans le Treeview
def afficher_donnees_treeview(treeview, df):
    logging.debug("Affichage des données dans le Treeview.")
    for row in treeview.get_children():
        treeview.delete(row)
    treeview["columns"] = list(df.columns)
    treeview.heading("#0", text="Index")
    for col in df.columns:
        treeview.heading(col, text=col)
    for i, row in df.iterrows():
        treeview.insert("", "end", text=i, values=list(row))


def verifier_changements():
    global fichier_xlsb, dernieres_valeurs, ws
    try:
        logging.debug(f"Vérification des changements dans le fichier XLSB : {fichier_xlsb}")
        nouvelles_valeurs = []
        row = 1
        while True:
            valeur_col_A = ws.Cells(row, 1).Value
            if valeur_col_A is None:
                break
            nouvelles_valeurs.append(valeur_col_A)
            if valeur_col_A not in dernieres_valeurs:
                dernieres_valeurs.append(valeur_col_A)
                # Comparaison avec la base de données manifeste
                if base_de_donnees_manifeste is not None:
                    for idx, row_bdd in base_de_donnees_manifeste.iterrows():
                        if row_bdd['AWB'] == valeur_col_A:
                            afficher_informations_correspondantes(row_bdd)
                            enregistrer_modification(valeur_col_A, row_bdd)
                            enregistrer_dans_excel(valeur_col_A, row_bdd)
                            trouver_info_adresse(valeur_col_A,row_bdd)
                            break
            row += 1
        
        dernieres_valeurs = nouvelles_valeurs
    except Exception as e:
        logging.error(f"Erreur lors de la vérification des changements du fichier XLSB : {e}")
    
    # Re-vérifier les changements après 2 secondes
    root.after(2000, verifier_changements)

    # Fonction pour enregistrer les modifications détectées
def enregistrer_modification(valeur_col_A, informations):
    try:
        with open("modifications_detectees.log", "a") as log_file:
            log_file.write(f"Modification détectée pour la valeur {valeur_col_A} : {informations.to_dict()}\n")
        logging.debug(f"Modification enregistrée pour la valeur {valeur_col_A} : {informations.to_dict()}")
    except Exception as e:
        logging.error(f"Erreur lors de l'enregistrement de la modification : {e}")


# Fonction pour enregistrer les données dans un fichier Excel
def enregistrer_dans_excel(valeur_col_A, informations):
    date_str = datetime.now().strftime("%d-%m-%Y")
    dossier = f"Dispatch_output_{date_str}"
    if not os.path.exists(dossier):
        os.makedirs(dossier)
    
    file_name =os.path.join(dossier,f"dispatch_{date_str}.xlsx") 
    dest_city = informations['DestCity']
    consignee_address = informations['ConsigneeAddress']
    
    # Parcourir les adresses et trouver les informations correspondantes
    adresse = consignee_address
     # Utiliser la nouvelle méthode pour trouver les informations d'adresse
    fokontany, kaomina, distrika, secteur, region, province = trouver_info_adresse(consignee_address, df_localisation)

    if os.path.exists(file_name):
        workbook = load_workbook(file_name)
    else:
        workbook = Workbook()
        workbook.remove(workbook.active)  # Supprimer la feuille par défaut

    if dest_city not in workbook.sheetnames:
        sheet = workbook.create_sheet(title=dest_city)
        # Ajouter les en-têtes uniquement si la feuille est nouvellement créée
        sheet.append(['Secteur', 'AWB', 'name', 'adresse','telephone'])
    else:
        sheet = workbook[dest_city]

    # Vérifier si l'AWB est déjà présent pour ce secteur
    awb_present = False
    for row in sheet.iter_rows(min_row=2, values_only=True):
        if row[0] == secteur and row[1] == informations['AWB']:
            awb_present = True
            break

    if not awb_present:
        # Ajouter les informations à la feuille correspondante
        sheet.append([
            secteur,
            informations['AWB'],
            informations['ConsigneeName'],
            informations['ConsigneeAddress'],
            informations['ConsigneeTel'],
            informations["DestCity"]
        ])

    workbook.save(file_name)
    logging.debug(f"Les données ont été enregistrées dans {file_name} sous l'onglet {dest_city}")

# Fonction pour créer le modèle Word
def creer_modele_word():
    date_str = datetime.now().strftime("%d-%m-%Y")
    dossier = f"Template{date_str}"
    if not os.path.exists(dossier) :
        os.makedirs(dossier)
    

    file_name = choisir_fichier_excel()
    if not file_name:
        logging.error("Aucun fichier sélectionné.")
        return
    
    word_file_name = f"Template_{datetime.now().strftime('%d-%m-%Y')}.docx"

    if not os.path.exists(file_name):
        logging.error(f"Le fichier Excel {file_name} n'existe pas.")
        return

    # Charger le fichier Excel
    workbook = load_workbook(file_name)
    
    document = Document()

    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        
        # Collecter les informations pour chaque secteur
        secteur_data = {}
        for row in sheet.iter_rows(min_row=2, values_only=True):
            telephone = row[4]
            adresse = row[3]
            secteur = row[0]
            name = row[2]  # Assuming name is in the third column
            if secteur not in secteur_data:
                secteur_data[secteur] = []
            secteur_data[secteur].append(name)
        
        # Créer une section pour chaque secteur
        for secteur, names in secteur_data.items():
            name_counts = Counter(names)
            
            # Ajouter une table avec 4 sections sur une page
            table = document.add_table(rows=2, cols=2)
            table.autofit = False
            widths = [Pt(480), Pt(480)]
            for i, width in enumerate(widths):
                for cell in table.columns[i].cells:
                    cell.width = width
            
            cell_index = 0
            for name, count in name_counts.items():
                cell = table.cell(cell_index // 2, cell_index % 2)
                paragraph = cell.paragraphs[0]
                # Ajout des informations de contact
                paragraph.add_run("EXP : CELERO MADAGASCAR\n").bold = True
                paragraph.add_run("ADRESSE: LOT IVF 4 FITROAFANA TALAMATY IVATO, ANTANANARIVO\n").bold = True
                paragraph.add_run("Téléphone : 0381942854\n").bold = True
                
                #information des destinataire
                run = paragraph.add_run(f"\n DEST: {name} ({count})\n")
                run.font.size = Pt(12)
                
                tex =paragraph.add_run(f"\nAdresse: {adresse} \n")
                tex.font.size = Pt(12)
                
                adres =paragraph.add_run(f"\nTelephone: {telephone} \n")
                adres.font.size = Pt(12)
               

                paragraph.alignment =  WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                cell_index += 1

                # S'il y a plus de 6 éléments, ajouter une nouvelle table
                if cell_index == 4:
                    document.add_page_break()
                    table = document.add_table(rows=2, cols=2)
                    table.autofit = False
                    for i, width in enumerate(widths):
                        for cell in table.columns[i].cells:
                            cell.width = width
                    cell_index = 0
    word_file_names = os.path.join(dossier, word_file_name)
    document.save(word_file_names)
    logging.debug(f"Le modèle Word a été créé et enregistré sous {word_file_names}")



# Fonction pour enregistrer les modifications détectées
def enregistrer_modification(valeur_col_A, informations):
    try:
        with open("modifications_detectees.log", "a") as log_file:
            log_file.write(f"Modification détectée pour la valeur {valeur_col_A} : {informations.to_dict()}\n")
        logging.debug(f"Modification enregistrée pour la valeur {valeur_col_A} : {informations.to_dict()}")
    except Exception as e:
        logging.error(f"Erreur lors de l'enregistrement de la modification : {e}")

# Créer une fenêtre Tkinter principale
root = tk.Tk()
root.title("Surveillance des Modifications XLSB")
root.geometry("1200x600")
root.config(bg=background)

# Frame pour contenir le Treeview
frame_treeview = ttk.Frame(root)
frame_treeview.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

# Création du Treeview
treeview = ttk.Treeview(frame_treeview)
treeview.pack(fill=tk.BOTH, expand=True)

scrollbary = ttk.Scrollbar(frame_treeview, orient=tk.VERTICAL, command=treeview.yview)
scrollbary.pack(side=tk.RIGHT, fill=tk.Y)

scrollbarx = ttk.Scrollbar(frame_treeview, orient=tk.HORIZONTAL, command=treeview.xview)
scrollbarx.pack(side=tk.BOTTOM, fill=tk.X)

treeview.configure(yscrollcommand=scrollbary.set, xscrollcommand=scrollbarx.set)
treeview["columns"] = ("AWB", "ConsigneeName", "ConsigneeAddress", "DestCity")
treeview.heading("#0", text="Index")
treeview.column("#0", stretch=tk.NO, width=50)
treeview.heading("AWB", text="AWB")
treeview.column("AWB", stretch=tk.NO, width=100)
treeview.heading("ConsigneeName", text="Consignee Name")
treeview.column("ConsigneeName", stretch=tk.NO, width=150)
treeview.heading("ConsigneeAddress", text="Consignee Address")
treeview.column("ConsigneeAddress", stretch=tk.NO, width=200)
treeview.heading("DestCity", text="Destination City")
treeview.column("DestCity", stretch=tk.NO, width=150)

# Bouton pour charger et afficher la base de données manifeste
btn_charger_base_de_donnees = ttk.Button(root, text="Choisir le fichier Excel pour la base de données manifeste", command=charger_et_afficher_base_de_donnees)
btn_charger_base_de_donnees.pack(pady=10)

#bouton pour creer un template pour template
btn_template = ttk.Button(root, text='creer template',command=creer_modele_word)
btn_template.pack(pady=10)



# Bouton pour quitter l'application
btn_quitter = ttk.Button(root, text="Quitter", command=root.destroy)
btn_quitter.pack(pady=10)

# Version corrigée sans l'option height
label_bottom = ttk.Label(root, text="E-mail: wzafitsara@gmail.com", width=10, anchor='e')

root.mainloop()
