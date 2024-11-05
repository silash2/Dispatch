import openpyxl
from tkinter import *
from tkinter import messagebox, filedialog
from tkinter.ttk import Combobox, Scrollbar, Treeview
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


background = "#06283D"

root = Tk()
root.title("Cession document")
root.geometry("700x620+700+100")
root.config(bg=background)
root.resizable(width=False, height=False)

# Variable globale pour le manifeste
base_de_donnee_manifest = None

# Télécharger et enregistrer le chemin de l'image
def get_image():
    global image_path    
    image_path = './data/Logo celero/celero_logo.png'
    if image_path:
        messagebox.showinfo('info', f'Image chargée avec succès depuis {image_path}')
    else:
        messagebox.showwarning('Avertissement', 'Aucune image sélectionnée')
    return image_path

# Recherche et charge le manifeste en question
def get_manifest():
    global base_de_donnee_manifest
    # Types de données
    file_types = [
        ("Excel files", "*.xls;*.xlsx"),
        ("All files", "*.*")
    ]
    
    file_path = filedialog.askopenfilename(filetypes=file_types)
    if file_path:
        base_de_donnee_manifest = pd.read_excel(file_path)
        messagebox.showinfo('info', f'Manifeste chargé avec succès depuis {file_path}')
    else:
        messagebox.showwarning('Avertissement', 'Aucun fichier sélectionné')
    return file_path, base_de_donnee_manifest

# Rechercher l'information dans le manifeste
def find_information():
    lta = lta_entry.get()
    info = awb_entry.get()
    weight = weight_entry.get()
    if base_de_donnee_manifest is not None:
        for idx, row_bdd in base_de_donnee_manifest.iterrows():
            if row_bdd['AWB'] == info:
                # Récupérer les informations nécessaires
                name = row_bdd['ConsigneeName']
                awb = row_bdd['AWB']
                pcs = row_bdd['PCS']
                currency = weight*8.39
                return name, awb, pcs,weight, currency,lta
    else:
        messagebox.showwarning('Avertissement', 'Aucun manifeste chargé')
    return None, None, None, None, None, None

# Fonction pour ajouter une bordure autour d'un paragraphe
def set_paragraph_border(paragraph):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    framePr = OxmlElement('w:pBdr')
    border = OxmlElement('w:top')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), '4')
    border.set(qn('w:space'), '1')
    border.set(qn('w:color'), '000000')
    framePr.append(border)
    
    border = OxmlElement('w:left')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), '4')
    border.set(qn('w:space'), '1')
    border.set(qn('w:color'), '000000')
    framePr.append(border)

    border = OxmlElement('w:bottom')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), '4')
    border.set(qn('w:space'), '1')
    border.set(qn('w:color'), '000000')
    framePr.append(border)
    
    border = OxmlElement('w:right')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), '4')
    border.set(qn('w:space'), '1')
    border.set(qn('w:color'), '000000')
    framePr.append(border)
    
    pPr.append(framePr)


#faire l'impression
def print_doc():
    att_fret()
    bon_a_delivrer()
    ordre_transit()
    messagebox.showinfo('info', 'cession document effectuer avec succes')

#Attestation de fret
def att_fret():
    date_str = datetime.now().strftime("%d-%m-%Y")
    name, awb, pcs,weight, currency, lta = find_information()
    if awb is None:
        messagebox.showwarning('Avertissement', 'Informations introuvables ou AWB incorrect')
        return
    
    date_st = datetime.now().strftime("%d/%m/%Y")
    doc = Document()

    # Dimensions de la page A4
    section = doc.sections[0]
    section.page_height = Pt(841.89)
    section.page_width = Pt(595.28)

    # Ajout de l'image en haut à droite
    if image_path:
        header = section.header
        paragraph = header.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(1.25))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    #ajout de la date
    dat = doc.add_paragraph()
    dat.add_run(f"Fait à Antananarivo le {date_st}.")
    dat.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    #ajout name
    conv = doc.add_paragraph()
    conv.add_run(f"{ name}")

    doc.add_paragraph("\n ")

    p = doc.add_paragraph()
    run = p.add_run("ATTESTATION DE FRET \n")
    run.bold = True
    run.font.size = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ajout des informations de contact
    p = doc.add_paragraph("\nNous soussignés CELERO MADAGASCAR Sarl attestons par la présente que le transport des colis dont les caractéristiques ci-dessous nous a été confié :")
    
    
    # Ajout des informations récupérées dans un cadre centré
    info_paragraph = doc.add_paragraph()
    run = info_paragraph.add_run(f"MAWB: {lta}\n")
    run.add_break()
    run = info_paragraph.add_run(f"HAWB: {awb}\n")
    run.add_break()
    run = info_paragraph.add_run(f"PCS: {pcs}\n")
    run.add_break()
    run = info_paragraph.add_run(f"Weight: {weight} Kg\n")
    run.add_break()
    run = info_paragraph.add_run("Nature du colis : Clothing or daily necessities").bold=True
    info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ajouter une bordure autour du paragraphe
    set_paragraph_border(info_paragraph)

    doc.add_paragraph("\n ")

    fn = doc.add_paragraph()
    fn.add_run(f"Le fret correspondant est de {currency} USD de HONG KONG à Madagascar.La présente est établie pour servir et valoir ce que de droit")
    doc.add_paragraph("\n ")

    dat = doc.add_paragraph()
    dat.add_run("Service Opérations \n").bold=True
    dat.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p.add_run("CELERO Madagascar").bold=True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    outputFile = f'Att_fret{date_str}.docx'
    doc.save(outputFile)
#bon a delivrer 
def bon_a_delivrer():

    date_str = datetime.now().strftime("%d-%m-%Y")
    name, awb, pcs,weight, currency, lta = find_information()
    if awb is None:
        messagebox.showwarning('Avertissement', 'Informations introuvables ou AWB incorrect')
        return
    
    date_st = datetime.now().strftime("%d/%m/%Y")
    doc = Document()

    # Dimensions de la page A4
    section = doc.sections[0]
    section.page_height = Pt(841.89)
    section.page_width = Pt(595.28)

    # Ajout de l'image en haut à droite
    if image_path:
        header = section.header
        paragraph = header.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(1.25))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p = doc.add_paragraph()
    run = p.add_run("Bon a Delivrer")
    run.bold = True
    run.font.size = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ajout des informations de contact
    p = doc.add_paragraph(f"\nNous vous prions de livrer les marchandises citées ci-après {name}sur présentation de la déclaration en douane :")
    doc.add_paragraph("\n ")

    # Ajout des informations récupérées dans un cadre centré
    info_paragraph = doc.add_paragraph()
    run = info_paragraph.add_run(f"MAWB: {lta}\n")
    run.add_break()
    run = info_paragraph.add_run(f"AWB: {awb}\n")
    run.add_break()
    run = info_paragraph.add_run(f"PCS: {pcs}\n")
    run.add_break()
    run = info_paragraph.add_run(f"Weight: {weight} Kg\n")
    run.add_break()
    run = info_paragraph.add_run("Nature du colis : Clothing or daily necessities").bold=True
    info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ajouter une bordure autour du paragraphe
    set_paragraph_border(info_paragraph)

    doc.add_paragraph("\n ")

    fn = doc.add_paragraph()
    fn.add_run(f"La présente est établie pour servir et valoir ce que de droit.")
    doc.add_paragraph("\n ")

    dat = doc.add_paragraph()
    dat.add_run("Service Opérations \n")
    dat.alignment = WD_ALIGN_PARAGRAPH.RIGHT
  
    p.add_run("CELERO Madagascar")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    

    outputFile = f'bon_a_delivrer{date_str}.docx'
    doc.save(outputFile)

def ordre_transit():
    date_str = datetime.now().strftime("%d-%m-%Y")
    name, awb, pcs,weight, currency, lta = find_information()
    if awb is None:
        messagebox.showwarning('Avertissement', 'Informations introuvables ou AWB incorrect')
        return
    
    date_st = datetime.now().strftime("%d/%m/%Y")
    doc = Document()

    # Dimensions de la page A4
    section = doc.sections[0]
    section.page_height = Pt(841.89)
    section.page_width = Pt(595.28)

    # Ajout de l'image en haut à droite
    if image_path:
        header = section.header
        paragraph = header.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(1.25))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p = doc.add_paragraph()
    run = p.add_run("ORDRE DE TRANSIT \nIMPORT")
    run.bold = True
    run.font.size = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ajout des informations de contact
    p = doc.add_paragraph("\nNous soussignés")
    p = doc.add_paragraph()
    p.add_run("Nom de la Société : CELERO Madagascar\n").bold = True
    p.add_run("Nom de contact : Sylvain\n").bold = True
    p.add_run("Adresse complète : lot IVF 4 Fitroafana Talatamaty Ivato\n").bold = True
    p.add_run("Téléphone : 0381942854\n").bold = True
    
    #ajout des convenances
    conv = doc.add_paragraph()
    conv.add_run(f"Donnons ordre par la présente à { name} la procédure douanière des colis suivants :\n")
    doc.add_paragraph("\n ")

    # Ajout des informations récupérées dans un cadre centré
    info_paragraph = doc.add_paragraph()
    run = info_paragraph.add_run(f"MAWB: {lta}\n")
    run.add_break()
    run = info_paragraph.add_run(f"AWB: {awb}\n")
    run.add_break()
    run = info_paragraph.add_run(f"PCS: {pcs}\n")
    run.add_break()
    run = info_paragraph.add_run(f"Weight: {weight} Kg\n")
    run.add_break()
    run = info_paragraph.add_run("Nature du colis : Clothing or daily necessities").bold=True
    info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ajouter une bordure autour du paragraphe
    set_paragraph_border(info_paragraph)

    doc.add_paragraph("\n ")
    #les convenances
    fn = doc.add_paragraph()
    fn.add_run(f"Cet ordre de transit import est établi pour servir et valoir ce que de droit.\n VEUILLEZ DELIVRER LE BAD à {name}")
    doc.add_paragraph("\n ")

    #date de l'edition
    dat = doc.add_paragraph()
    dat.add_run(f"Fait à Antananarivo le {date_st}.")
    dat.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    outputFile = f'ordre_transit{date_str}.docx'
    doc.save(outputFile)

# Entête de l'application
Label(root, text="CELERO EXPRESS CESSION DOCUMENT", width=10, height=3, bg="white", fg='blue', font="arial 20 bold").pack(side=TOP, fill=X)
# Bas de l'application
Label(root, text="E-mail: wzafitsara@gmail.com", width=10, height=2, anchor='e').pack(side=BOTTOM, fill=X)

awb = StringVar()
Label(root, text="Entrer l'AWB du colis", bg=background, foreground="white", font='arial 10').place(x=280, y=120)
awb_entry = Entry(root, textvariable=awb, width=30)
awb_entry.place(x=250, y=150)
lta = StringVar()
Label(root, text="Entrer LTA ", bg=background, foreground="white", font='arial 10').place(x=310, y=170)
lta_entry = Entry(root, textvariable=lta, width=30)
lta_entry.place(x=250, y=200)

weight = float()
Label(root, text="Entrer le poids a l'eclatement", bg=background, foreground="white", font="arial 10 ").place(x=250, y=230)
weight_entry = Entry(root, textvariable=weight,width=30)
weight_entry.place(x=250, y=250)

# Bouton pour télécharger un fichier
Button(root, text="Télécharger Manifest", font='arial 12', width=30, command=get_manifest).place(x=200, y=300)
# Bouton pour préparer les documents
Button(root, text="Générer la cession doc", font='arial 12', width=30, command=print_doc).place(x=200, y=350)

# Bouton retour
Button(root, text="Retour", width=30, font='arial 12').place(x=200, y=400)
# Bouton quitter
Button(root, text="Quitter", width=30, font='arial 12', foreground="red", command=root.destroy).place(x=200, y=450)

root.mainloop()
